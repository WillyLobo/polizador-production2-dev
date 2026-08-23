"""
Construcción del cuerpo de las resoluciones de viáticos (solicitud Chaco,
solicitud exterior, incorporación) sobre el documento base único: el
`.docx` vigente en `EncabezadoDocumento` (subido vía `ActualizarEncabezado`),
que aporta header/footer/sección/márgenes y tiene el cuerpo vacío.

Reemplaza al esquema anterior de 3 plantillas estáticas con tags Jinja de
docxtpl. Acá el cuerpo se arma párrafo a párrafo con `python-docx`,
replicando el formato exacto (fuente, espaciados, sangrías y tabuladores
con línea de puntos) que tenían esas plantillas, para que el documento
generado se vea igual.

El punto que motivó este módulo: el tabulador con línea de puntos que
alinea el monto a la derecha en el Artículo 2° (uno por agente comisionado)
es, en el .docx original, una propiedad de párrafo de OOXML
(`w:tabs`/`w:leader="dot"`), no un carácter "\\t" en el texto. Como ya no
hay una plantilla de la que heredar esa propiedad, se recrea acá con
`paragraph_format.tab_stops` + `run.add_tab()`.
"""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Pt

FONT_NAME = "Arial"
FONT_SIZE = Pt(12)

# Posición del tab derecho del Artículo 2°: w:pos="9921" (dxa/twips) en las
# plantillas originales. 1pt = 20 twips.
ARTICULO_DOS_TAB_POS = Pt(9921 / 20)
# Sangría de primera línea de los "considerandos": w:ind firstLine="708".
CONSIDERANDO_INDENT = Pt(708 / 20)

# Espacio debajo del subpárrafo de cada agente del Artículo 2°, para separarlos
# visualmente entre sí cuando hay varios.
ARTICULO_DOS_ESPACIO_ENTRE_AGENTES = 120


def _set_font(run, bold=False):
	run.font.name = FONT_NAME
	run.font.size = FONT_SIZE
	run.bold = bold
	# python-docx solo fija w:rFonts ascii/hAnsi vía font.name; las
	# plantillas originales también fijan eastAsia/cs al mismo valor.
	rpr = run._element.get_or_add_rPr()
	rfonts = rpr.find(qn("w:rFonts"))
	if rfonts is not None:
		rfonts.set(qn("w:eastAsia"), FONT_NAME)
		rfonts.set(qn("w:cs"), FONT_NAME)


def _add_run(paragraph, text, bold=False):
	run = paragraph.add_run(text)
	_set_font(run, bold=bold)
	return run


def _add_paragraph(doc, *, jc=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=276, first_line_indent=None, keep_with_next=False):
	p = doc.add_paragraph()
	pf = p.paragraph_format
	pf.alignment = jc
	pf.space_before = Pt(before / 20)
	pf.space_after = Pt(after / 20)
	if line is not None:
		pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
		pf.line_spacing = line / 240
	if first_line_indent is not None:
		pf.first_line_indent = first_line_indent
	# Evita que Word corte la página justo entre este párrafo y el
	# siguiente (p. ej. "EL PRESIDENTE..." separado de "RESUELVE:").
	pf.keep_with_next = keep_with_next
	return p


def _add_articulo_dos_tabs(paragraph):
	tab_stops = paragraph.paragraph_format.tab_stops
	tab_stops.add_tab_stop(Pt(0), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
	tab_stops.add_tab_stop(ARTICULO_DOS_TAB_POS, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def add_heading_inline(doc, label, texto, *, before=240, after=0):
	"""Párrafo tipo 'VISTO: <texto>' o 'Artículo 1º: <texto>': etiqueta en
	negrita seguida del texto en formato normal, dentro del mismo párrafo."""
	p = _add_paragraph(doc, before=before, after=after)
	_add_run(p, label, bold=True)
	_add_run(p, texto, bold=False)
	return p


def add_heading_bloque(doc, texto, *, jc=WD_ALIGN_PARAGRAPH.JUSTIFY, before=240, after=0, line=276, keep_with_next=False):
	"""Párrafo íntegramente en negrita: 'CONSIDERANDO:', 'RESUELVE:', etc."""
	p = _add_paragraph(doc, jc=jc, before=before, after=after, line=line, keep_with_next=keep_with_next)
	_add_run(p, texto, bold=True)
	return p


def add_considerando(doc, texto):
	p = _add_paragraph(doc, before=240, first_line_indent=CONSIDERANDO_INDENT)
	_add_run(p, texto)
	return p


def add_articulo_dos_fila(doc, fila):
	"""Una fila de agente del Artículo 2°: dos párrafos —
	'nombre_cuil [tab-con-puntos] $monto' y el subpárrafo de detalle—,
	ambos con los mismos tab-stops de línea de puntos que la plantilla
	original."""
	p_monto = _add_paragraph(doc, before=0, after=0)
	_add_articulo_dos_tabs(p_monto)
	_add_run(p_monto, fila["nombre_cuil"])
	tab_run = p_monto.add_run()
	_set_font(tab_run)
	tab_run.add_tab()
	_add_run(p_monto, f"${fila['monto']}")

	p_sub = _add_paragraph(doc, before=0, after=ARTICULO_DOS_ESPACIO_ENTRE_AGENTES)
	_add_articulo_dos_tabs(p_sub)
	_add_run(p_sub, fila["subparrafo"])
	return p_monto, p_sub


def build_resolucion_docx(base_docx_file, *, visto_texto, parrafos, incluir_ultimo_parrafo,
						   articulo_uno, articulo_dos_filas, articulo_tres, articulo_cuatro,
						   articulo_cinco, considerandos_fijos_finales):
	"""
	Arma el cuerpo completo (VISTO/CONSIDERANDO/RESUELVE/Artículos) sobre
	`base_docx_file` (path o file-like: el .docx vigente de
	EncabezadoDocumento, con header/footer/sección y cuerpo vacío) y
	devuelve un BytesIO con el .docx final.

	`parrafos` son los considerandos editables (parrafo_uno..N); si
	`incluir_ultimo_parrafo` es False se omite el último (mismo
	comportamiento que el `{%if actuacion.solicitud_dia_inhabil%}` de las
	plantillas originales). `considerandos_fijos_finales` son los
	considerandos de boilerplate que siguen siempre a los editables.
	"""
	doc = Document(base_docx_file)

	add_heading_inline(doc, "VISTO: ", visto_texto)
	add_heading_bloque(doc, "CONSIDERANDO:")

	parrafos_a_incluir = parrafos if incluir_ultimo_parrafo else parrafos[:-1]
	for parrafo in parrafos_a_incluir:
		add_considerando(doc, parrafo)
	for parrafo in considerandos_fijos_finales:
		add_considerando(doc, parrafo)

	add_heading_bloque(doc, "EL PRESIDENTE DEL INSTITUTO PROVINCIAL DE DESARROLLO URBANO Y VIVIENDA", jc=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
	add_heading_bloque(doc, "RESUELVE:", jc=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=None)

	add_heading_inline(doc, "Artículo 1º: ", articulo_uno, before=240, after=240)
	add_heading_inline(doc, "Artículo 2º: ", "Anticipese por la Dirección Contable de este Instituto, lo que se enuncia:", before=0, after=240)
	for fila in articulo_dos_filas:
		add_articulo_dos_fila(doc, fila)
	add_heading_inline(doc, "Artículo 3°: ", articulo_tres, before=240, after=0)
	add_heading_inline(doc, "Artículo 4º: ", articulo_cuatro, before=240, after=0)
	add_heading_inline(doc, "Artículo 5º: ", articulo_cinco, before=240, after=0)

	out = io.BytesIO()
	doc.save(out)
	out.seek(0)
	return out
